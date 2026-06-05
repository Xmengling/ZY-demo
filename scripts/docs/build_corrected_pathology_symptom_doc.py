from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT = PROJECT_ROOT / "tcm_rag_demo" / "data" / "经方病理分类症状归属整理表.docx"


PATHOLOGY_ROWS = [
    ("表证", "表实", "表闭、邪在表，人体欲从汗解而汗不得出。", "恶寒重、发热、无汗、头项强痛、身疼痛、骨节疼痛、脉浮紧、鼻塞、咳喘、体表闭塞感。", "浮肿、身重、烦躁、呕逆、气上冲、痉挛、皮肤痒。", "表实重在无汗、脉浮紧、身疼。若同时痰饮浮肿、小便不利，要兼看水实。"),
    ("表证", "表虚", "肌表功能虚弱，营卫不和，表不固而易汗出恶风。", "汗出、恶风、怕风、发热或低热、脉浮缓、自汗、容易感冒、皮肤营养不良、久病表不解。", "关节痛、黄汗、盗汗、浮肿、麻痹不仁、疮口不收。", "表虚不是里虚。表虚重点在汗出恶风、体表敏感；里虚重点在胃弱、食少、乏力、脉软。"),
    ("里证", "里实", "里部有实邪、积滞、结实、硬满，常见腹满便难或心下腹部抗拒。", "腹满、腹硬、心下痞硬、按之抗手、便秘、大便不通、腹痛拒按、肠中结滞。", "胀满、烦躁、谵语、潮热、少腹急结、痰食宿积。", "里实常与里热、气实、血实同见。仅腹胀不一定里实，食后胀而胃弱可属里虚。"),
    ("里证", "里虚", "中焦胃气不足，吸收转化弱，津液气血来源不足。", "纳差、食少、吃一点就胀、胃弱、乏力、但欲寐、脉软弱、舌淡嫩、便溏、容易呕、久病虚弱。", "汗后虚、产后虚、腹中急痛、口渴但水不化、恢复慢。", "里虚是很多虚证的根。水虚、血虚、气虚可由里虚发展，但不能全部等同于里虚。"),
    ("里证", "里寒", "里部寒盛或虚寒，阳气不足，胃肠功能偏冷偏衰。", "四肢厥冷、腹痛喜温、胃冷、呕吐清水、下利清谷、便溏、口不渴或喜热饮、脉沉迟微、舌淡水滑。", "肠鸣、食冷即不适、咳喘寒饮、尿清长、精神萎靡。", "怕冷要分表证恶寒和里寒。表证多脉浮头项强痛；里寒多腹冷、下利、脉沉微。"),
    ("里证", "里热", "里部热盛、热结或热迫津液外越。", "恶热、身热、高热、口渴欲冷饮、汗出而热、汗黏、心烦、尿黄、大便干、便秘、舌红苔黄、脉数或洪。", "口腔溃疡、牙疼、黄疸、脓涕、痰黄、皮肤红肿热痛、蹬被子、谵语。", "里热与半热都可见热象。里热偏全身热盛、渴、大便干；半热偏口苦、胸胁、孔窍、往来寒热。"),
    ("半证", "半热", "半表半里层面的热郁，多见胸胁、孔窍、情志、节律问题。", "口苦、咽干、目眩、胸胁苦满、往来寒热、心烦、易怒、多梦易醒、鼻炎、咽炎、耳鸣、黄涕脓涕。", "恶心欲吐、食少、盗汗、子时发作、口腔溃疡、牙疼、目赤、咳痰黄。", "半热不是单纯里热。抓胸胁、孔窍、情志、节律；若见大便干、恶热、渴饮明显，再合看里热。"),
    ("水证", "水实", "水饮、痰饮、湿、水毒停留或排泄不利。", "痰多、咳喘、短气、清嗓子、鼻涕多、浮肿、面肿、眼肿、下肢肿、小便不利、心悸、眩晕、苔水滑或腻、脉滑或沉滑。", "身重、关节痛、黄疸、腹中水声、恶心呕吐、睡觉打呼噜、喝水少或水不化、腹满。", "水实可造成口渴或不渴。停饮重时水不化津，外表像缺水，根子仍可有水实。"),
    ("水证", "水虚", "可用水虚理解津液不足、津液不生或组织失濡。", "口渴、咽干、唇干、皮肤干、大便干、小便少、舌干少津、烦渴、消渴、干咳少痰。", "虚热、筋脉拘急、眼干、心烦失眠、热病后期虚弱。", "口渴不一定水虚，也可能里热或水实水不化津。看少津、干燥、汗吐下后、久病耗伤。"),
    ("血证", "血实", "瘀血、血分实邪、血行不畅或血热瘀结。", "舌紫暗、舌底静脉怒张、唇暗、固定刺痛、少腹硬满、腹中包块、结节、癥瘕、痛经血块、皮肤斑疹。", "健忘、狂躁、产后腹痛、跌打瘀肿、胸胁刺痛、出血夹瘀。", "血实抓固定、暗、硬、块、久、刺痛。胀痛走窜多偏气实，重着肿胀多偏水实。"),
    ("血证", "血虚", "血液营养不足，筋脉、肌肉、心神失养。", "面色淡、唇淡、头晕、心悸、失眠、麻木、抽筋、筋急、月经量少、闭经、产后虚、脉细。", "腹痛拘急、眼干、皮肤干、健忘、虚劳、爪甲不荣。", "血虚与水虚都可见干、虚、失养。血虚偏血色、月经、筋脉、心神；水虚偏口渴、少津、大便干。"),
    ("气证", "气实", "气机壅滞、上冲、逆满、不降。", "胸闷、胸满、胸胁满、腹胀、胃胀、心下满、嗳气、打嗝、咽中如有物、气上冲、奔豚、咳逆上气、鼓音。", "头面冲胀、心悸、烦躁、胁痛、胸痛、呕逆、短气。", "气实常与水实并见。气实抓胀满冲逆；水实抓痰饮苔滑、小便不利、浮肿。"),
    ("气证", "气虚", "推动、升降、固摄和恢复力不足，常与里虚相连。", "乏力、气短、懒言、动则喘、声音低、汗出、易疲劳、脉弱、食少、久病虚弱。", "下陷感、便溏、脱肛、子宫下垂、恢复慢、反复感冒。", "气虚资料中常并入里虚或表虚。若以胃弱食少为主归里虚；以汗出恶风为主归表虚；以气短乏力推动不足为主归气虚。"),
    ("阴证", "阴证", "整体机能沉衰、寒虚重、反应低下，是较深层的虚寒或衰败状态。", "精神萎靡、但欲寐、四肢厥逆、脉微细欲绝、下利清谷、畏寒明显、面色晦暗、声音低弱、反应迟钝。", "冷汗、亡阳样表现、久病体弱、老人虚寒、病后误汗误下后转弱。", "阴证不是单个症状，而是整体趋势。若局部有热象但人整体弱、脉微、厥冷，也要警惕寒热错杂或阴证转机。"),
]


OVERLAP_ROWS = [
    ("口渴", "里热、水虚、水实", "里热渴多喜冷饮；水虚渴伴口咽干、舌少津；水实可因水不化津而渴，常伴苔滑、小便不利。"),
    ("怕冷/恶寒", "表实、表虚、里寒、阴证", "表实多无汗脉浮紧；表虚多汗出恶风；里寒多腹冷下利；阴证看精神萎靡、脉微、厥冷。"),
    ("汗出", "表虚、里热、水虚、气虚、阴证", "汗出恶风归表虚；热迫汗出归里热；汗后伤津归水虚；动则汗出归气虚；冷汗肢厥看阴证。"),
    ("发热", "表实、表虚、半热、里热", "表证发热常兼恶寒脉浮；半热有往来寒热、胸胁口苦；里热恶热渴饮便干。"),
    ("咳喘/短气", "水实、气实、表实、里寒、气虚", "痰饮水逆归水实；上冲逆满归气实；表闭肺代偿归表实；寒饮归里寒；动则喘乏力归气虚。"),
    ("胸胁满", "半热、气实、水实", "口苦往来寒热偏半热；胀满冲逆偏气实；伴痰饮眩悸苔滑偏水实。"),
    ("腹胀", "气实、里虚、里寒、里实", "胀满嗳气偏气实；食少食后胀偏里虚；冷痛便溏偏里寒；便秘硬满拒按偏里实。"),
    ("便秘", "里实、里热、气实、血实、水虚", "硬满拒按归里实；热结归里热；胀满不通归气实；少腹瘀痛归血实；干燥少津归水虚。"),
    ("鼻炎/鼻塞/脓涕", "表实、表虚、半热、水实、里热", "鼻塞新感可表；反复孔窍炎症多半热；涕多痰多苔滑多水实；黄脓热象明显可兼里热。"),
    ("心悸", "水实、半热、血虚、气实、气虚", "停饮上扰可悸；烦热扰心可悸；血不养心可悸；气上冲可悸；气虚推动弱也可悸。"),
    ("失眠多梦", "半热、血虚、水虚、气虚", "烦躁口苦多半热；心血不足多血虚；虚热口干多水虚；疲劳心悸气短多气虚。"),
    ("舌淡水滑", "里寒、里虚、水实", "淡为虚寒，水滑为水饮，食少脉软可兼里虚。"),
    ("舌红苔黄", "里热、半热", "全身热盛、便干、渴饮偏里热；口苦咽干、胸胁孔窍偏半热。"),
    ("舌底静脉怒张/唇暗", "血实", "典型瘀血抓手，可与气实、水实、里热同见，但血实证据权重较高。"),
]


def add_run(paragraph, text, bold=False, size=8.5, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            cell = row.cells[i]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))


def setup(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 15, "1F4D78"), ("Heading 2", 12, "2E74B5")]:
        st = doc.styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    add_run(p, "整理口径：", True, 9, "1F4D78")
    add_run(p, text, False, 9)


def table(doc, headers, rows, widths, size=7.3):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, True, 8, "0B2545")
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            margins(cells[i], 70, 90, 70, 90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            if i < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, val, True, size, "1F4D78" if i == 1 else "0B2545")
            else:
                add_run(p, val, False, size)
    set_widths(tbl, widths)
    return tbl


def main():
    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "经方病理分类下的症状归属整理", True, 18, "0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "按用户指定分类重整：表证、里证、半证、水证、血证、气证、阴证", False, 9, "555555")

    note(doc, "本版不再使用“十二字病理”标题，而按你指定的分类：表证分表实/表虚，里证分里实/里虚/里寒/里热，半证保留半热，水证分水实/水虚，血证分血实/血虚，气证分气实/气虚，另列阴证。一个症状仍允许归入多个病理。")

    doc.add_heading("一、病理分类与症状归属总表", level=1)
    table(
        doc,
        ["大类", "病理", "核心理解", "常见主症", "可兼见症状", "鉴别要点"],
        PATHOLOGY_ROWS,
        [1.5, 1.6, 4.4, 6.2, 5.1, 6.6],
        7.0,
    )

    doc.add_page_break()
    doc.add_heading("二、同一症状的多病理归属与追问", level=1)
    note(doc, "这一表更适合放进智能问诊逻辑：症状先进入候选病理，再用追问、舌脉、二便、寒热、饮水、胸腹等信息加权。")
    table(
        doc,
        ["症状/体征", "可能相关病理", "追问与鉴别"],
        OVERLAP_ROWS,
        [4.0, 5.0, 15.2],
        8.0,
    )

    doc.add_heading("三、程序化建议", level=1)
    suggestions = [
        ("保留层级", "数据库建议存 pathology_group 和 pathology_label 两列，例如 group=表证，label=表实。"),
        ("一症多标", "症状不要单选病理，允许一个 answer 产生多个 candidate_labels，并附 evidence_weight。"),
        ("追问优先", "当候选标签冲突时生成追问，例如怕冷要区分表实、表虚、里寒、阴证。"),
        ("保留原词", "资料里的表证、表实、表虚、水虚、气虚、阴证等原词要保留，不要强行映射成另一套。"),
        ("输出证据链", "最终分析要显示“症状 -> 候选病理 -> 追问修正 -> 最终病理”的链路。"),
    ]
    table(doc, ["建议", "说明"], suggestions, [4.6, 19.6], 8.2)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Codex 根据用户指定病理分类重新整理。", False, 8, "666666")
    doc.save(OUT)


if __name__ == "__main__":
    main()
