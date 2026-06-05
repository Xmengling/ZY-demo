from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT = PROJECT_ROOT / "docs" / "pathology" / "十二字病理症状归属整理表.docx"


PATHOLOGY_ROWS = [
    {
        "label": "半表",
        "core": "人体想从表解、从外排邪或恢复肌表功能；原资料多写作表证、表实、表虚。",
        "main": "恶寒、恶风、发热、头项强痛、身疼痛、脉浮、鼻塞、喷嚏、皮肤痒、关节疼、表证咳喘。",
        "secondary": "汗出、无汗、气上冲、呕逆、身重、浮肿、小便不利、鼻炎反复。",
        "distinguish": "半表不是体表局部病，而是患病后呈现出的症状反应。表虚多汗出恶风、脉缓；表实多无汗、身疼、脉浮紧。若同时有痰饮浮肿、小便不利，要合看水实。",
        "sources": "《伤寒论》讲稿中反复解释太阳病为脉浮、头项强痛、恶寒等症状反应；药物集萃中桂枝、麻黄、生姜、防风等多归表证相关标签。",
    },
    {
        "label": "半热",
        "core": "半表半里或孔窍、胸胁、情志层面的热郁；常与少阳、咽口鼻耳目、烦躁睡眠相关。",
        "main": "口苦、咽干、目眩、胸胁苦满、往来寒热、心烦、易怒、多梦易醒、口腔溃疡、牙疼、咽炎、鼻炎、黄涕脓涕、目赤。",
        "secondary": "恶心欲吐、食少、耳鸣、盗汗、子时发作、咳嗽痰黄、皮肤疮疡、妇科带下阴痒伴热象。",
        "distinguish": "半热和里热都可见热象。半热更偏胸胁、孔窍、节律、情志和往来寒热；里热更偏全身实热、口渴、大汗、大便干、舌红苔黄。",
        "sources": "小柴胡、柴胡桂枝等讲解中多见胸胁苦满、口苦、烦、呕、往来寒热；药物集萃中柴胡、黄芩、黄连、菊花、百合等常归半热。",
    },
    {
        "label": "半虚",
        "core": "半表半里层面的虚性不安、收敛不足或神志不宁；常需龙骨、牡蛎、柏子仁、山萸肉等线索帮助判断。",
        "main": "虚烦、惊悸、心悸、失眠、多梦、精神恍惚、易惊、烦躁而虚、盗汗、自汗、遗精、遗尿。",
        "secondary": "头晕、耳鸣、腰膝酸软、汗出不止、虚性亢奋、久病后心神不安。",
        "distinguish": "半虚不是单纯里虚。里虚偏胃气与中焦功能弱；半虚偏神志、收敛、节律和虚性躁动。若有热烦，常与半热并见。",
        "sources": "药物集萃中龙骨、牡蛎、柏子仁、山萸肉、朱砂、鳖甲等出现半虚标签，相关药性多指向安神、收敛、虚劳、强壮。",
    },
    {
        "label": "里热",
        "core": "里部热盛或热结，热刺激强，常压过恶寒反应，表现为恶热、渴、烦、大便干等。",
        "main": "怕热、恶热、身热、高热、口渴欲饮、汗出而热、汗黏、心烦、谵语、口舌生疮、牙疼、尿黄、大便干、便秘、舌红、苔黄、脉数或洪。",
        "secondary": "黄疸、脓涕、痰黄、出血鲜红、皮肤红肿热痛、蹬被子、烦躁不眠。",
        "distinguish": "里热常有热迫津液外出或耗津，所以可同时见汗多、口渴、津虚。若只见口苦咽干胸胁满，未必是里热，可能偏半热。",
        "sources": "讲稿中阳明、白虎汤、白虎加人参汤等讲解强调里热不恶寒、渴、身热；药物集萃中石膏、知母、大黄、芒硝、茵陈等归里热。",
    },
    {
        "label": "里寒",
        "core": "里部阳气不足或寒饮内停，功能沉衰，常见冷、利、痛、呕、脉微沉迟。",
        "main": "怕冷、四肢厥冷、腹痛喜温、胃冷、呕吐清水、下利清谷、便溏、口不渴或喜热饮、精神萎靡、脉沉迟微、舌淡水滑。",
        "secondary": "咳喘属寒饮、心下痞硬偏冷、腹中冷痛、肠鸣、食冷即不适、尿清长。",
        "distinguish": "里寒常与里虚、水实并见。单见怕冷不一定是里寒，半表恶寒也怕冷；需看是否有胃肠冷利、脉沉微、舌淡水滑。",
        "sources": "附子、干姜、吴茱萸、蜀椒、乌头等药物集萃均指向里寒、水实、里虚；讲稿中四逆、理中、吴茱萸类方多围绕阴寒、呕利、厥冷。",
    },
    {
        "label": "里虚",
        "core": "胃气、中焦、吸收转化功能不足，是津液、气血生成和恢复力不足的基础。",
        "main": "纳差、食少、吃一点就胀、胃弱、乏力、但欲寐、便溏、久病虚弱、脉软弱、舌淡嫩、容易呕、恢复慢。",
        "secondary": "口渴但水不化、津液不足、汗后虚、产后虚、腹中急痛、心下痞而不实。",
        "distinguish": "里虚不是所有虚证。它更偏胃肠吸收和中焦生化。津虚、血虚可由里虚导致，但不等同于里虚。",
        "sources": "人参讲解中强调健胃生津、改善胃对水分吸收；大枣、甘草、粳米、山药、白术等多归里虚。",
    },
    {
        "label": "水实",
        "core": "水饮、痰饮、湿、水毒停滞或排泄不利；可在表、半、里各层出现。",
        "main": "痰多、咳喘、短气、清嗓子、鼻涕多、浮肿、面肿、眼肿、下肢肿、小便不利、心悸、眩晕、身重、苔水滑或腻、脉滑或沉滑。",
        "secondary": "关节疼痛、皮肤水肿、黄疸、腹中水声、恶心呕吐、睡觉打呼噜、喝水少或水不化、腹满。",
        "distinguish": "水实可造成口渴或不渴：停饮重时水不化津，外表看缺水，根子仍可有水实。需结合苔滑、脉滑、小便、痰饮、浮肿。",
        "sources": "药物集萃中麻黄、半夏、茯苓、猪苓、泽泻、葶苈子、杏仁等大量归水实；讲稿多次用停饮、水毒解释咳喘、短气、悸、小便不利。",
    },
    {
        "label": "血虚",
        "core": "血液、营养、滋养不足，筋脉肌肉心神失养，常与津虚、里虚互相牵连。",
        "main": "面色淡、唇淡、头晕、心悸、失眠、麻木、抽筋、筋急、月经量少、闭经、产后虚、爪甲不荣、脉细。",
        "secondary": "腹痛拘急、虚性疼痛、眼干、皮肤干、健忘、虚劳、肌肉萎弱。",
        "distinguish": "血虚和津虚都可干、虚、失养。血虚偏血色、月经、筋脉、心神；津虚偏口渴、咽干、舌少津、大便干。",
        "sources": "药物集萃中当归、芍药、川芎、阿胶、鸡子黄、柏子仁等归血虚；芍药、当归类方多对应腹痛拘急、血不养筋。",
    },
    {
        "label": "气实",
        "core": "气机壅滞、上冲、逆满、不降，常表现为胀、满、冲、逆、咳喘。",
        "main": "胸闷、胸满、胸胁满、腹胀、胃胀、心下满、嗳气、打嗝、咽中如有物、气上冲、奔豚、咳逆上气、短气、鼓音。",
        "secondary": "头面冲胀、心悸、烦躁、胁痛、胸痛、呕逆、便秘伴胀满。",
        "distinguish": "气实常与水实并见，水饮上逆也会短气、咳喘；气实抓“胀满冲逆”，水实抓“痰饮苔滑小便浮肿”。",
        "sources": "桂枝、厚朴、枳实、陈皮、杏仁、旋覆花等药物集萃多归气实；桂枝讲解反复提到气上冲、奔豚、诸逆。",
    },
    {
        "label": "血实",
        "core": "瘀血、血分实邪、血行不畅或血热瘀结，常见固定、暗紫、结块、刺痛。",
        "main": "舌紫暗、舌底静脉怒张、唇暗、固定刺痛、少腹硬满、腹中包块、结节、癥瘕、经血有块、痛经、皮肤斑疹、健忘、狂躁。",
        "secondary": "出血夹瘀、黄疸伴瘀、产后腹痛、跌打瘀肿、胸胁痛、局部红紫热痛。",
        "distinguish": "血实不等于所有疼痛。固定、暗、硬、块、久、刺痛是抓手；气实疼痛多胀满走窜，水实疼痛多重着肿胀。",
        "sources": "桃仁、红花、牡丹皮、茜草、蒲黄、虻虫、蛴螬等药物集萃归血实；AI医案资料中舌底静脉怒张明确归瘀血。",
    },
    {
        "label": "津虚",
        "core": "津液不足、阴液枯燥、组织失濡，常由热耗、汗吐下后、久病或胃吸收差导致。",
        "main": "口渴、咽干、唇干、皮肤干、大便干、小便少、舌红少津、舌干、烦渴、消渴、虚热、干咳少痰。",
        "secondary": "筋脉拘急、心烦失眠、眼干、便秘、盗汗后虚、热病后期虚弱。",
        "distinguish": "口渴不一定就是津虚：里热可渴，水实水不化津也可渴。津虚重点看干燥、少津、耗伤、虚热和大便干。",
        "sources": "人参、麦冬、天花粉、百合、玉竹、生地等讲解常涉及生津、止渴、津液枯燥；讲稿中特别强调津液虚与停饮可同时出现。",
    },
    {
        "label": "津实",
        "core": "资料中很少直接以“津实”标注。可暂理解为津液输布壅滞、黏滞不化的状态，实际操作多并入水实或痰饮观察。",
        "main": "痰黏、涎多、口中黏腻、便黏、苔腻、分泌物黏稠、胸咽黏阻感。",
        "secondary": "口不渴或渴不欲饮、清嗓子、鼻涕黏、湿热黄浊分泌物。",
        "distinguish": "若是水液停留、浮肿、小便不利，优先归水实；若是干燥少津，归津虚；津实建议作为程序里的辅助标签，不宜替代水实。",
        "sources": "十二字体系里保留津实，但本文件夹资料主要使用水实、水虚/津虚；因此此项为结合体系后的补充整理。",
    },
]


OVERLAP_ROWS = [
    ("口渴", "里热、津虚、水实", "热盛可渴，津亏可渴，停饮水不化津也可渴。看舌津、大便、小便、苔滑与否。"),
    ("发热", "半表、半热、里热", "半表多恶寒脉浮；半热多往来寒热胸胁孔窍；里热多恶热口渴汗出。"),
    ("怕冷/恶寒", "半表、里寒", "半表恶寒多伴脉浮头项强痛；里寒怕冷多伴下利、腹冷、脉沉微。"),
    ("汗出", "半表、里热、津虚、半虚", "表虚可汗出恶风；里热可热迫汗出；汗后伤津则津虚；虚性不固可半虚。"),
    ("头痛", "半表、半热、气实、血实", "表证头项强痛；半热可目眩头痛；气上冲可头面胀；瘀血多固定刺痛。"),
    ("咳喘/短气", "水实、气实、半表、里寒", "痰饮水逆、气机不降、表闭肺代偿、寒饮上逆都可咳喘。"),
    ("胸胁满", "半热、气实、水实", "少阳胸胁苦满偏半热；胀满冲逆偏气实；伴痰饮悸眩苔滑偏水实。"),
    ("腹胀", "气实、里虚、里寒、里实", "气滞可胀，胃弱食后胀属里虚，寒凝腹胀属里寒，便秘硬满属里实。"),
    ("便秘", "里热、气实、血实、津虚", "热结、气滞、瘀结、津亏肠燥都可便秘，需看热象、胀满、瘀象、干燥程度。"),
    ("鼻炎/鼻塞/脓涕", "半表、半热、水实、里热", "鼻塞可表；脓涕黄涕多热；涕多痰多苔滑多水实；反复孔窍炎症多半热。"),
    ("心悸", "水实、半虚、血虚、气实", "停饮上扰可悸，虚性不安可悸，血不养心可悸，气上冲也可悸。"),
    ("失眠多梦", "半热、半虚、血虚、津虚", "烦热扰心、虚性不安、血不养心、津亏虚热都可影响睡眠。"),
    ("舌淡水滑", "里寒、里虚、水实", "淡为虚寒，水滑为水饮，若胃弱食少则兼里虚。"),
    ("舌红苔黄", "里热、半热", "全身热盛便干口渴偏里热；口苦咽干胸胁孔窍偏半热。"),
    ("舌底静脉怒张/唇暗", "血实", "典型瘀血抓手，也可与气实、水实、热结并见。"),
]


def add_run(paragraph, text, bold=False, size=9, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            cell = row.cells[i]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))


def setup_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    add_run(p, "整理原则：", True, 9, "1F4D78")
    add_run(p, text, False, 9)


def add_table(doc, headers, rows, widths, font_size=7.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, True, 8, "0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i], 70, 90, 70, 90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, val, True, font_size, "1F4D78")
            else:
                add_run(p, val, False, font_size)
    set_table_widths(table, widths)
    return table


def main():
    doc = Document()
    setup_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "十二字病理下的常见症状归属整理", True, 18, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "依据 tcm_rag_demo/data 内胡希恕、李冠杰讲稿、药物集萃、方剂思维导图与医案资料整理", False, 8.5, "555555")

    add_note_box(
        doc,
        "一个症状不固定归属一个病理。这里采用“病理核心 + 主症 + 兼症 + 鉴别”的方式整理，尤其注意半表/里热/半热、水实/津虚、气实/水实、血虚/血实之间的交叉。",
    )

    doc.add_heading("一、十二字病理症状总表", level=1)
    rows = [
        (r["label"], r["core"], r["main"], r["secondary"], r["distinguish"])
        for r in PATHOLOGY_ROWS
    ]
    add_table(
        doc,
        ["病理", "核心理解", "常见主症", "可兼见症状", "鉴别要点"],
        rows,
        [1.6, 5.0, 6.4, 5.2, 7.2],
        7.2,
    )

    doc.add_page_break()
    doc.add_heading("二、常见症状多归属矩阵", level=1)
    add_note_box(
        doc,
        "这一页用于程序设计和问诊追问：当一个症状同时指向多个病理时，不要直接定性，应结合舌脉、二便、寒热、饮水、胸腹、病程等信息追问。",
    )
    add_table(
        doc,
        ["症状/体征", "可能相关病理", "追问与鉴别"],
        OVERLAP_ROWS,
        [4.0, 5.2, 15.0],
        8,
    )

    doc.add_heading("三、资料学习后的实现建议", level=1)
    suggestions = [
        ("症状字段不要单选病理", "同一症状允许打到多个标签，例如口渴同时候选里热、津虚、水实，再由追问和舌脉降权或升权。"),
        ("病理打分要保留证据", "每个标签后面要保存原始症状证据，如“舌淡水滑→里寒/水实”，“胸胁苦满+口苦→半热”。"),
        ("追问围绕分歧点", "口渴时追问饮水量、喜冷喜热、小便、舌津；怕冷时追问是否脉浮头痛或腹冷下利。"),
        ("先分层再分寒热虚实", "先判断半表、半热、里，再判断热寒虚实、水血气津。这样比直接按现代病名推荐方剂更稳。"),
        ("保留老师原体系词", "资料里有表实、表虚、水虚、气虚、阴证、半寒等词，程序可保留为原始标签，再映射到十二字体系。"),
    ]
    add_table(doc, ["建议", "说明"], suggestions, [5.0, 19.2], 8)

    doc.add_heading("四、参考资料范围", level=1)
    refs = [
        "《伤寒论-上册》胡希恕、李冠杰讲稿合订本-原文标注.docx / .txt",
        "《伤寒论-下册》胡希恕、李冠杰讲稿合订本-原文标注.docx",
        "《金匮要略》胡希恕、李冠杰讲稿合订本-条文标注.docx",
        "李冠杰经方病理辨证体系经方药物集萃（目录)(1).docx",
        "伤寒金匮方剂思维导图.reparsed.txt",
        "AI学习医案.docx",
    ]
    for ref in refs:
        p = doc.add_paragraph(style=None)
        add_run(p, "• ", False, 8.5, "1F4D78")
        add_run(p, ref, False, 8.5)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Codex 整理：用于中医智能问诊病理标签与追问逻辑设计。", False, 8, "666666")
    doc.save(OUT)


if __name__ == "__main__":
    main()
