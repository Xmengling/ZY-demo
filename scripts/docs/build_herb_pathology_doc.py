from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE_KEY = "\u7ecf\u65b9\u75c5\u7406\u8fa8\u8bc1\u4f53\u7cfb"
TARGET_LABELS = [
    "\u534a\u8868",
    "\u534a\u70ed",
    "\u534a\u865a",
    "\u91cc\u70ed",
    "\u91cc\u5bd2",
    "\u91cc\u865a",
    "\u6c34\u5b9e",
    "\u8840\u865a",
    "\u6c14\u5b9e",
    "\u8840\u5b9e",
    "\u6d25\u865a",
    "\u6d25\u5b9e",
]

LABEL_NORMALIZE = {
    "\u6c34\u865a": "\u6d25\u865a",
    "\u6c14\u865a": "\u91cc\u865a",
    "\u8868\u8bc1": "\u534a\u8868",
    "\u8868\u5b9e": "\u534a\u8868",
    "\u8868\u865a": "\u534a\u8868",
}

LABEL_GRAB = {
    "\u534a\u8868": "\u89e3\u8868/\u808c\u8868\u95ed\u963b\u6216\u8868\u865a\u4e0d\u56fa",
    "\u534a\u70ed": "\u6e05\u534a\u8868\u534a\u91cc\u70ed/\u5b54\u7a8d\u70ed\u6bd2",
    "\u534a\u865a": "\u6536\u655b\u5b89\u795e/\u865a\u6027\u6447\u52a8\u6216\u865a\u70e6",
    "\u91cc\u70ed": "\u6e05\u91cc\u70ed/\u70e6\u6e34\u3001\u9ec4\u75b8\u3001\u4fbf\u79d8\u70ed\u7ed3",
    "\u91cc\u5bd2": "\u6e29\u91cc\u6563\u5bd2/\u5455\u5229\u3001\u8179\u75db\u3001\u56db\u9006",
    "\u91cc\u865a": "\u5065\u80c3\u8865\u4e2d/\u751f\u6d25\u6db2\u3001\u590d\u80c3\u6c14",
    "\u6c34\u5b9e": "\u53bb\u6c34\u996e/\u75f0\u996e\u3001\u6d6e\u80bf\u3001\u5c0f\u4fbf\u4e0d\u5229",
    "\u8840\u865a": "\u517b\u8840\u8865\u865a/\u8840\u4e0d\u8db3\u3001\u7b4b\u8109\u5931\u517b",
    "\u6c14\u5b9e": "\u964d\u9006\u7834\u6c14/\u80f8\u6ee1\u8179\u80c0\u3001\u54b3\u9006\u4e0a\u6c14",
    "\u8840\u5b9e": "\u6d3b\u8840\u7834\u7600/\u7600\u8840\u3001\u75c7\u7626\u3001\u75bc\u75db",
    "\u6d25\u865a": "\u751f\u6d25\u6da6\u71e5/\u6e34\u3001\u865a\u70ed\u3001\u6d25\u6db2\u4e0d\u8db3",
    "\u6d25\u5b9e": "\u6d25\u6db2\u58c5\u6ede/\u9700\u7ed3\u5408\u539f\u8bc1\u8fa8\u522b",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))


def parse_rows(source: Path) -> list[dict[str, str]]:
    src_doc = Document(str(source))
    rows = []
    for idx, para in enumerate(src_doc.paragraphs):
        text = para.text.strip()
        if not text or SOURCE_KEY not in text:
            continue
        raw = text.split("\uff1a", 1)[1] if "\uff1a" in text else text.split(":", 1)[-1]
        raw = raw.strip().rstrip("\u3002.")
        if "\uff1a" in raw:
            herb, labels = raw.split("\uff1a", 1)
        elif "\uff0c" in raw:
            herb, labels = raw.split("\uff0c", 1)
        elif "," in raw:
            herb, labels = raw.split(",", 1)
        else:
            continue
        herb = herb.strip()
        labels = labels.strip().replace("\uff0c", " ")
        labels = re.sub(r"\s+", " ", labels)
        label_list = [x.strip() for x in labels.split(" ") if x.strip()]
        normalized = []
        extras = []
        for label in label_list:
            mapped = LABEL_NORMALIZE.get(label, label)
            if mapped in TARGET_LABELS:
                if mapped not in normalized:
                    normalized.append(mapped)
            else:
                if label not in extras:
                    extras.append(label)
        grab = "\uff1b".join(LABEL_GRAB[x] for x in normalized if x in LABEL_GRAB)
        if extras:
            grab = (grab + "\uff1b" if grab else "") + "\u8865\u5145\u6807\u7b7e\u9700\u56de\u5230\u539f\u6587\u8fa8\u522b\uff1a" + "\u3001".join(extras)
        rows.append(
            {
                "seq": str(len(rows) + 1),
                "herb": herb,
                "source_labels": " ".join(label_list),
                "normalized_labels": " ".join(normalized) if normalized else "\u5f85\u786e\u8ba4",
                "extra_labels": " ".join(extras) if extras else "",
                "grab": grab,
            }
        )
    return rows


def add_run(paragraph, text: str, bold=False, size=9, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [("Heading 1", 16, "1F4D78"), ("Heading 2", 12, "2E74B5")]:
        st = styles[name]
        st.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(5)


def build_doc(source: Path, out: Path) -> None:
    rows = parse_rows(source)
    label_counter = Counter()
    for row in rows:
        label_counter.update(row["normalized_labels"].split())

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "\u5e38\u89c1\u7ecf\u65b9\u836f\u7269\u4e0e\u5341\u4e8c\u5b57\u75c5\u7406\u6807\u7b7e\u6620\u5c04\u8868", True, 18, "0B2545")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "\u6765\u6e90\uff1a", True, 9, "555555")
    add_run(meta, source.name + "\uff1b\u62bd\u53d6\u4f9d\u636e\uff1a\u201c\u7ecf\u65b9\u75c5\u7406\u8fa8\u8bc1\u4f53\u7cfb\uff1a\u201d\u539f\u6587\u6807\u6ce8", False, 9, "555555")

    note_table = doc.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    note_cell = note_table.cell(0, 0)
    set_cell_shading(note_cell, "F4F6F9")
    set_cell_margins(note_cell, 120, 160, 120, 160)
    p = note_cell.paragraphs[0]
    add_run(p, "\u8bf4\u660e\uff1a", True, 9, "1F4D78")
    add_run(
        p,
        "\u672c\u8868\u4fdd\u7559\u539f\u8d44\u6599\u6807\u7b7e\uff0c\u5e76\u6309\u4f60\u5e38\u7528\u7684\u5341\u4e8c\u5b57\u75c5\u7406\u4f53\u7cfb\u505a\u4e86\u5f52\u7eb3\u3002\u539f\u6587\u4e2d\u7684\u201c\u8868\u8bc1/\u8868\u5b9e/\u8868\u865a\u201d\u6682\u5f52\u5165\u201c\u534a\u8868\u201d\uff0c\u201c\u6c34\u865a\u201d\u5f52\u5165\u201c\u6d25\u865a\u201d\uff0c\u201c\u6c14\u865a\u201d\u5f52\u5165\u201c\u91cc\u865a\u201d\uff1b\u5176\u4ed6\u4e0d\u5728\u5341\u4e8c\u6807\u7b7e\u5185\u7684\u8bcd\u653e\u5165\u201c\u8865\u5145/\u5f85\u786e\u8ba4\u6807\u7b7e\u201d\u3002",
        False,
        9,
    )

    doc.add_heading("\u4e00\u3001\u6807\u7b7e\u901f\u67e5", level=1)
    summary = doc.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["\u75c5\u7406\u6807\u7b7e", "\u672c\u8868\u51fa\u73b0\u6b21\u6570", "\u7528\u836f\u6293\u624b"]
    for i, h in enumerate(headers):
        cell = summary.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        add_run(cell.paragraphs[0], h, True, 9, "0B2545")
    for label in TARGET_LABELS:
        cells = summary.add_row().cells
        vals = [label, str(label_counter.get(label, 0)), LABEL_GRAB.get(label, "")]
        for i, val in enumerate(vals):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            add_run(cells[i].paragraphs[0], val, False, 8.5)
    set_table_width(summary, [3.0, 2.4, 18.0])

    doc.add_heading("\u4e8c\u3001\u836f\u7269\u6620\u5c04\u8868", level=1)
    main = doc.add_table(rows=1, cols=6)
    main.style = "Table Grid"
    main.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["\u5e8f", "\u836f\u7269", "\u539f\u6587\u75c5\u7406\u6807\u7b7e", "\u5f52\u5165\u5341\u4e8c\u5b57\u6807\u7b7e", "\u8865\u5145/\u5f85\u786e\u8ba4\u6807\u7b7e", "\u4e34\u5e8a\u6293\u624b\u6458\u8981"]
    for i, h in enumerate(headers):
        cell = main.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        add_run(cell.paragraphs[0], h, True, 8, "0B2545")

    for row in rows:
        cells = main.add_row().cells
        vals = [
            row["seq"],
            row["herb"],
            row["source_labels"],
            row["normalized_labels"],
            row["extra_labels"],
            row["grab"],
        ]
        for i, val in enumerate(vals):
            set_cell_margins(cells[i], 70, 90, 70, 90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cells[i].paragraphs[0]
            if i in (0, 1):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(para, val, False, 7.2)

    set_table_width(main, [1.0, 2.0, 4.0, 4.2, 3.2, 10.0])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "\u7531 Codex \u6839\u636e\u539f\u6587\u6807\u6ce8\u62bd\u53d6\u6574\u7406\uff0c\u5efa\u8bae\u4f5c\u4e3a\u5b66\u4e60\u7b14\u8bb0\u548c\u540e\u7eed\u7a0b\u5e8f\u5316\u914d\u7f6e\u7684\u521d\u7a3f\u3002", False, 8, "666666")

    doc.save(str(out))


if __name__ == "__main__":
    source = Path(sys.argv[1])
    out = Path(sys.argv[2])
    build_doc(source, out)
