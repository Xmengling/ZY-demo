# -*- coding: utf-8 -*-
"""问诊医案批量导出 Word（标准版：原文附件 + 结构化病例摘要）。"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from ..models import ConsultModuleHint, ConsultSession, ConsultSymptomPreset
from .consult_attachments import (
    IMAGE_EXTS,
    get_ext,
    read_attachment_bytes,
)
from .consult_case_summary import (
    build_consult_summary_groups,
    format_consult_summary_groups,
    merge_session_intake,
)

STATUS_LABELS = {
    "collecting": "采集中",
    "analyzed": "已分析",
    "completed": "已完成",
}

IMAGE_WIDTH = Cm(16)
MAX_EMBED_BYTES = 20 * 1024 * 1024


class ConsultCaseExportError(Exception):
    pass


def _loads_json(raw: str) -> dict[str, Any]:
    import json

    if not raw:
        return {}
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError:
        return {}


def _loads_list(raw: str) -> list[Any]:
    import json

    if not raw:
        return []
    try:
        data = json.loads(raw)
        return data if isinstance(data, list) else []
    except json.JSONDecodeError:
        return []


def load_symptom_sections(db: Session) -> list[dict[str, Any]]:
    rows = (
        db.query(ConsultSymptomPreset)
        .order_by(
            ConsultSymptomPreset.module_order,
            ConsultSymptomPreset.block_order,
            ConsultSymptomPreset.id,
        )
        .all()
    )
    sections: dict[str, dict[str, Any]] = {}
    for row in rows:
        section = sections.get(row.module_key)
        if not section:
            section = {
                "key": row.module_key,
                "order": row.module_order,
                "title": row.module_title,
                "blocks": [],
            }
            sections[row.module_key] = section
        section["blocks"].append(
            {
                "label": row.block_label,
                "symptoms": _loads_list(row.symptoms),
            }
        )
    return list(sections.values())


def _set_run_font(run, *, size_pt: float = 10.5, bold: bool = False, color: RGBColor | None = None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(document: Document, text: str, *, bold: bool = False, size_pt: float = 10.5, color: RGBColor | None = None):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph


def _chief_complaint_preview(intake: dict[str, Any], fallback: str = "") -> str:
    text = str(intake.get("chief_complaint") or fallback or "").strip()
    text = re.sub(r"\s+", " ", text)
    return text[:40] + ("…" if len(text) > 40 else "")


def _status_label(status: str) -> str:
    return STATUS_LABELS.get(status or "", "采集中")


def _format_filter_line(chief: str, patient: str, doctor: str) -> str:
    parts = []
    if chief:
        parts.append(f"主诉含「{chief}」")
    if patient:
        parts.append(f"患者含「{patient}」")
    if doctor:
        parts.append(f"主诊医生含「{doctor}」")
    return "；".join(parts) if parts else "全部医案"


def _extract_text_attachment(raw: bytes, filename: str) -> str:
    ext = get_ext(filename)
    if ext in {".txt", ".md", ".csv"}:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return raw.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace").strip()

    if ext == ".docx":
        try:
            source = Document(io.BytesIO(raw))
            lines = [para.text.strip() for para in source.paragraphs if para.text.strip()]
            return "\n".join(lines).strip()
        except Exception:
            return ""

    if ext == ".pdf":
        try:
            import fitz

            doc = fitz.open(stream=raw, filetype="pdf")
            chunks = [page.get_text().strip() for page in doc if page.get_text().strip()]
            doc.close()
            return "\n\n".join(chunks).strip()
        except Exception:
            return ""
    return ""


def _add_attachment_section(document: Document, session: ConsultSession, intake: dict[str, Any], case_text: str):
    attachments = list(intake.get("attachments") or [])
    _add_paragraph(document, "一、医案原文（附件）", bold=True, size_pt=12)

    if attachments:
        for item in attachments:
            attachment_id = str(item.get("id") or "").strip()
            name = str(item.get("name") or "附件").strip()

            if not attachment_id:
                _add_paragraph(document, "（附件记录缺少 ID，无法读取文件）", size_pt=9, color=RGBColor(152, 162, 179))
                continue

            payload = read_attachment_bytes(session.id, attachment_id)
            if not payload:
                _add_paragraph(document, "（附件文件不存在或已被删除）", size_pt=9, color=RGBColor(152, 162, 179))
                continue

            raw, filename = payload
            ext = get_ext(filename or name)

            if ext in IMAGE_EXTS:
                if len(raw) > MAX_EMBED_BYTES:
                    _add_paragraph(
                        document,
                        f"（图片过大，未嵌入：{name}，{len(raw) // (1024 * 1024)}MB）",
                        size_pt=9,
                        color=RGBColor(152, 162, 179),
                    )
                    continue
                try:
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(io.BytesIO(raw), width=IMAGE_WIDTH)
                except Exception:
                    _add_paragraph(document, f"（无法嵌入图片：{name}）", size_pt=9, color=RGBColor(152, 162, 179))
                continue

            text = _extract_text_attachment(raw, filename or name)
            if text:
                for line in text.splitlines():
                    if line.strip():
                        _add_paragraph(document, line.rstrip())
                continue

            _add_paragraph(
                document,
                f"（原文文件：{name}，请在系统中查看附件）",
                size_pt=9,
                color=RGBColor(152, 162, 179),
            )
        return

    fallback = str(case_text or "").strip()
    if fallback:
        _add_paragraph(document, "（无附件，以下为系统存档原文）", size_pt=9, color=RGBColor(102, 112, 133))
        for line in fallback.splitlines():
            if line.strip():
                _add_paragraph(document, line.rstrip())
        return

    _add_paragraph(document, "（无原文附件）", size_pt=9, color=RGBColor(152, 162, 179))


def _add_structured_section(document: Document, intake: dict[str, Any], sections: list[dict[str, Any]]):
    _add_paragraph(document, "二、结构化医案", bold=True, size_pt=12)
    groups = build_consult_summary_groups(intake, sections)
    body = format_consult_summary_groups(groups)
    if not body.strip():
        _add_paragraph(document, "（暂无可导出的结构化内容）", size_pt=9, color=RGBColor(152, 162, 179))
        return

    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        heading = lines[0]
        if heading.startswith("【") and heading.endswith("】"):
            _add_paragraph(document, heading.strip("【】"), bold=True, size_pt=11)
            for line in lines[1:]:
                if not line.strip():
                    continue
                if "：" in line:
                    label, rest = line.split("：", 1)
                    paragraph = document.add_paragraph()
                    label_run = paragraph.add_run(f"{label}：")
                    _set_run_font(label_run, bold=True)
                    text_run = paragraph.add_run(rest)
                    _set_run_font(text_run)
                    paragraph.paragraph_format.line_spacing = 1.5
                else:
                    _add_paragraph(document, line)
        else:
            _add_paragraph(document, block)


def _add_case_header(document: Document, index: int, total: int, session: ConsultSession, intake: dict[str, Any]):
    patient = str(intake.get("patient_name") or session.patient_name or "—").strip() or "—"
    preview = _chief_complaint_preview(intake, session.title)
    title = f"医案 {index} / {total} · {patient}"
    if preview:
        title = f"{title} · {preview}"
    _add_paragraph(document, title, bold=True, size_pt=14)

    created = session.created_at
    if created.tzinfo is None:
        created = created.replace(tzinfo=timezone.utc)
    created_local = created.astimezone().strftime("%Y-%m-%d %H:%M")
    doctor = str(intake.get("doctor") or "—").strip() or "—"
    meta = f"创建时间：{created_local}    主诊医生：{doctor}    状态：{_status_label(session.status)}"
    _add_paragraph(document, meta, size_pt=9, color=RGBColor(102, 112, 133))


def build_cases_word(
    sessions: list[ConsultSession],
    sections: list[dict[str, Any]],
    *,
    chief_complaint: str = "",
    patient_name: str = "",
    doctor: str = "",
) -> bytes:
    if not sessions:
        raise ConsultCaseExportError("当前没有可导出的医案")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(10.5)

    now = datetime.now().astimezone()
    total = len(sessions)
    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run("医案记录导出")
    _set_run_font(cover_run, size_pt=18, bold=True)
    document.add_paragraph()

    _add_paragraph(document, f"导出时间：{now.strftime('%Y年%m月%d日 %H:%M')}", size_pt=10)
    _add_paragraph(document, f"导出条数：{total} 条", size_pt=10)
    _add_paragraph(
        document,
        f"筛选条件：{_format_filter_line(chief_complaint, patient_name, doctor)}",
        size_pt=10,
    )
    document.add_page_break()

    for index, session in enumerate(sessions, start=1):
        intake_data = _loads_json(session.intake_data)
        intake = merge_session_intake(session, intake_data)
        _add_case_header(document, index, total, session, intake)
        document.add_paragraph()
        _add_attachment_section(document, session, intake_data, session.case_text or "")
        document.add_paragraph()
        _add_structured_section(document, intake, sections)
        if index < total:
            document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_export_filename(count: int, now: datetime | None = None) -> str:
    current = now or datetime.now()
    stamp = current.strftime("%Y%m%d_%H%M")
    return f"医案导出_{stamp}_共{count}条_含原文.docx"
